"""
Generate sample CV files (DOCX + PDF) for testing the CV Matcher.
Run:  python create_samples.py
"""
from pathlib import Path
from docx import Document


def add_heading(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def add_section(doc: Document, title: str, content: str):
    add_heading(doc, title, level=2)
    doc.add_paragraph(content)


OUTPUT_DIR = Path("sample_cvs")
OUTPUT_DIR.mkdir(exist_ok=True)


# ── Sample 1 – Senior Python Developer ────────────────────────────────────────
def create_sample_1():
    doc = Document()
    doc.add_heading("Alice Johnson", 0)
    doc.add_paragraph("alice.johnson@email.com  |  +1-555-0101  |  linkedin.com/in/alicejohnson")

    add_section(doc, "Summary",
        "Senior Python Developer with 7 years of experience building scalable web applications "
        "and data pipelines. Passionate about clean code, RESTful APIs, and cloud infrastructure."
    )

    add_heading(doc, "Skills", 2)
    for s in ["Python, FastAPI, Flask, Django",
              "PostgreSQL, MySQL, MongoDB, Redis",
              "Docker, Kubernetes, AWS, CI/CD",
              "React, TypeScript, REST API, GraphQL",
              "Machine Learning, Pandas, NumPy, Scikit-learn",
              "Git, GitHub, Agile, Scrum"]:
        doc.add_paragraph(s, style="List Bullet")

    add_heading(doc, "Experience", 2)
    doc.add_paragraph("Senior Python Developer – TechCorp Inc.  (2019 – Present)")
    doc.add_paragraph(
        "• Led backend development of a microservices platform serving 2M users\n"
        "• Built ETL pipelines with Apache Airflow and PostgreSQL\n"
        "• Containerised services with Docker/Kubernetes on AWS"
    )
    doc.add_paragraph("Python Developer – DataSoft Ltd.  (2017 – 2019)")
    doc.add_paragraph(
        "• Developed Flask REST APIs and React frontends\n"
        "• Implemented machine learning models with scikit-learn\n"
        "• Wrote unit/integration tests achieving 90%+ coverage"
    )

    add_section(doc, "Education", "Bachelor of Science in Computer Science\nState University, 2017")

    path = OUTPUT_DIR / "alice_johnson_senior_python_dev.docx"
    doc.save(path)
    print(f"Created: {path}")
    return doc


# ── Sample 2 – Data Scientist ──────────────────────────────────────────────────
def create_sample_2():
    doc = Document()
    doc.add_heading("Bob Martinez", 0)
    doc.add_paragraph("bob.martinez@datamail.com  |  +1-555-0202")

    add_section(doc, "Professional Summary",
        "Data Scientist with 4 years of experience in machine learning, NLP, and data analysis. "
        "Proven track record of delivering actionable insights from complex datasets."
    )

    add_heading(doc, "Technical Skills", 2)
    for s in ["Python, R, SQL",
              "TensorFlow, PyTorch, Keras, Scikit-learn",
              "Pandas, NumPy, Matplotlib, Seaborn, Plotly",
              "PostgreSQL, MongoDB, Spark, Hadoop",
              "AWS SageMaker, Docker, Git",
              "NLP, Deep Learning, Computer Vision"]:
        doc.add_paragraph(s, style="List Bullet")

    add_heading(doc, "Work Experience", 2)
    doc.add_paragraph("Data Scientist – InsightAI  (2021 – Present)")
    doc.add_paragraph(
        "• Built NLP pipeline for sentiment analysis using transformers (HuggingFace)\n"
        "• Trained deep learning models achieving 94% accuracy on image classification\n"
        "• Presented findings to C-suite via Power BI dashboards"
    )
    doc.add_paragraph("Junior Data Analyst – Analytics Co.  (2020 – 2021)")
    doc.add_paragraph(
        "• Performed EDA and statistical modelling in Python and R\n"
        "• Automated reporting with pandas and matplotlib"
    )

    add_section(doc, "Education",
        "Master of Science in Data Science – Tech University, 2020\n"
        "Bachelor of Science in Mathematics – State College, 2018"
    )

    path = OUTPUT_DIR / "bob_martinez_data_scientist.docx"
    doc.save(path)
    print(f"Created: {path}")


# ── Sample 3 – Junior Full-Stack Developer ─────────────────────────────────────
def create_sample_3():
    doc = Document()
    doc.add_heading("Carol Smith", 0)
    doc.add_paragraph("carol.smith@webdev.io  |  +44-7700-900123")

    add_section(doc, "Objective",
        "Enthusiastic Junior Full-Stack Developer with 1 year of experience eager to grow "
        "in a collaborative team environment."
    )

    add_heading(doc, "Skills", 2)
    for s in ["JavaScript, TypeScript, React, Vue.js",
              "Node.js, Express, REST API",
              "HTML, CSS, Tailwind CSS, Bootstrap",
              "MySQL, MongoDB", "Git, GitHub, Docker (basic)", "Agile"]:
        doc.add_paragraph(s, style="List Bullet")

    add_heading(doc, "Experience", 2)
    doc.add_paragraph("Junior Developer – StartupXYZ  (2023 – Present)")
    doc.add_paragraph(
        "• Built React components and integrated REST APIs\n"
        "• Developed Node.js/Express microservices\n"
        "• Participated in daily standups and sprint planning"
    )

    add_section(doc, "Education",
        "Bachelor of Engineering in Software Engineering\nCity University, 2023"
    )

    path = OUTPUT_DIR / "carol_smith_junior_fullstack.docx"
    doc.save(path)
    print(f"Created: {path}")


# ── Sample 4 – DevOps Engineer ─────────────────────────────────────────────────
def create_sample_4():
    doc = Document()
    doc.add_heading("David Kim", 0)
    doc.add_paragraph("david.kim@devops.net  |  +82-10-5555-1234")

    add_section(doc, "Profile",
        "DevOps Engineer with 5 years of experience designing and maintaining CI/CD pipelines, "
        "cloud infrastructure, and container orchestration platforms."
    )

    add_heading(doc, "Core Skills", 2)
    for s in ["Docker, Kubernetes, Helm",
              "AWS, Azure, GCP, Terraform, Ansible",
              "Jenkins, GitHub Actions, GitLab CI, CircleCI",
              "Python, Bash, PowerShell",
              "Linux, Nginx, Apache",
              "Prometheus, Grafana, ELK Stack",
              "PostgreSQL, Redis", "Git, Agile, CI/CD"]:
        doc.add_paragraph(s, style="List Bullet")

    add_heading(doc, "Experience", 2)
    doc.add_paragraph("Senior DevOps Engineer – CloudBase  (2021 – Present)")
    doc.add_paragraph(
        "• Architected Kubernetes clusters on AWS EKS serving 50+ microservices\n"
        "• Reduced deployment time by 70% via optimised GitHub Actions pipelines\n"
        "• Managed infrastructure-as-code with Terraform across 3 regions"
    )
    doc.add_paragraph("DevOps Engineer – WebHost Ltd.  (2019 – 2021)")
    doc.add_paragraph(
        "• Built Jenkins CI/CD pipelines for 20+ projects\n"
        "• Containerised legacy applications with Docker\n"
        "• Automated server provisioning with Ansible"
    )

    add_section(doc, "Education",
        "Bachelor of Science in Computer Science\nKorea Tech University, 2019"
    )

    path = OUTPUT_DIR / "david_kim_devops_engineer.docx"
    doc.save(path)
    print(f"Created: {path}")


# ── Sample 5 – PDF CV ──────────────────────────────────────────────────────────
def create_sample_pdf():
    """Create a sample PDF CV using reportlab if available, else skip."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.enums import TA_LEFT, TA_CENTER

        path = OUTPUT_DIR / "emma_wilson_backend_dev.pdf"
        doc = SimpleDocTemplate(str(path), pagesize=A4,
                                topMargin=1.5*cm, bottomMargin=1.5*cm,
                                leftMargin=2*cm, rightMargin=2*cm)
        styles = getSampleStyleSheet()
        name_style   = ParagraphStyle("name",   fontSize=20, fontName="Helvetica-Bold", alignment=TA_CENTER)
        contact_style= ParagraphStyle("contact",fontSize=10, fontName="Helvetica",      alignment=TA_CENTER, textColor=colors.grey)
        h2_style     = ParagraphStyle("h2",     fontSize=13, fontName="Helvetica-Bold", spaceBefore=10, spaceAfter=4)
        body_style   = ParagraphStyle("body",   fontSize=10, fontName="Helvetica",      spaceAfter=4, leading=14)

        story = [
            Paragraph("Emma Wilson", name_style),
            Paragraph("emma.wilson@techdev.com  |  +1-555-0303  |  github.com/emmawilson", contact_style),
            HRFlowable(width="100%", thickness=1, color=colors.HexColor("#4F46E5"), spaceAfter=8),

            Paragraph("Summary", h2_style),
            Paragraph(
                "Backend Python Developer with 3 years of experience specialising in REST APIs, "
                "microservices architecture, and cloud-native applications on AWS.",
                body_style),

            Paragraph("Skills", h2_style),
            Paragraph("Python • FastAPI • Flask • Django • PostgreSQL • Redis • Docker • AWS • "
                      "Git • REST API • SQLAlchemy • Celery • Linux • CI/CD • Agile", body_style),

            Paragraph("Experience", h2_style),
            Paragraph("<b>Backend Developer – FinTech Solutions  (2022 – Present)</b>", body_style),
            Paragraph(
                "• Built high-throughput payment processing APIs with FastAPI and PostgreSQL<br/>"
                "• Implemented async task queues using Celery and Redis<br/>"
                "• Deployed services to AWS ECS with Docker containers", body_style),
            Paragraph("<b>Junior Python Developer – DevAgency  (2021 – 2022)</b>", body_style),
            Paragraph(
                "• Developed Django REST APIs for e-commerce platforms<br/>"
                "• Wrote unit tests with pytest achieving 85% coverage", body_style),

            Paragraph("Education", h2_style),
            Paragraph("Bachelor of Science in Computer Science – Northern University, 2021", body_style),
        ]

        doc.build(story)
        print(f"Created: {path}")
    except ImportError:
        print("reportlab not installed – skipping PDF sample. Run: pip install reportlab")


if __name__ == "__main__":
    create_sample_1()
    create_sample_2()
    create_sample_3()
    create_sample_4()
    create_sample_pdf()
    print("\nAll sample CVs created in ./sample_cvs/")
